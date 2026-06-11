"""
多格式文档解析器

支持格式：
  PDF（含扫描件 OCR 回退）/ Word (.docx) / Markdown / 纯文本

返回 list[ParsedPage]：每项保留 content / page_number / section_title / source，
供 chunker.py 做结构感知分块，最终写入 document_chunks 表的对应字段，
是 Phase 4 引用对齐（citation.py）的溯源依据。

OCR 路径依赖：
  - 系统安装 tesseract-ocr（含 chi_sim + eng 语言包）
  - Python 包 pymupdf（pip install pymupdf），仅用于 PDF → 图像渲染
  若 PyMuPDF 未安装，文本层不足的页面会记录 warning 并跳过，其余页面正常解析。
"""

import io
import re
import unicodedata
from collections import Counter
from collections.abc import Callable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pypdf
import pytesseract
import structlog
from markdown_it import MarkdownIt
from PIL import Image

logger = structlog.get_logger()

# PyMuPDF（fitz）：PDF 页面渲染为图像，可选但推荐安装
try:
    import fitz as _pymupdf  # type: ignore[import-untyped]
    _HAS_PYMUPDF: bool = True
except ImportError:
    _HAS_PYMUPDF = False


# ──────────────────────────────────────────────────────────────
# 数据模型
# ──────────────────────────────────────────────────────────────

@dataclass
class ParsedPage:
    """解析后的最小逻辑单元（PDF 一页 / Word 一节 / Markdown 一节 / 文本若干段）。

    Attributes:
        content:       清洗后的正文文本
        page_number:   原始页码；Markdown / txt 无物理页码时为 None
        section_title: 检测到的最近上层标题（用于分块时的结构感知断点）
        source:        原始文件名（引用溯源：chunk.source → 前端 CitationCard）
        metadata:      格式特有元信息（如 total_pages）
    """
    content: str
    page_number: int | None
    section_title: str | None
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)


# ──────────────────────────────────────────────────────────────
# 文本清洗工具
# ──────────────────────────────────────────────────────────────

_CTRL_CHARS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_EXCESS_NEWLINES = re.compile(r"\n{3,}")


def _clean_text(text: str) -> str:
    """基础清洗：NFC 规范化 → 去控制字符 → 压缩多余空行。"""
    text = unicodedata.normalize("NFC", text)
    text = _CTRL_CHARS.sub("", text)
    text = _EXCESS_NEWLINES.sub("\n\n", text)
    return text.strip()


def _detect_noisy_lines(raw_pages: list[str], threshold: float = 0.6) -> set[str]:
    """
    启发式页眉/页脚检测：

    对每页前 3 行与后 3 行进行统计，若某行在超过 threshold 比例的页面中出现，
    判定为重复噪声行（页眉/页脚/水印）并返回。

    设计取舍：
      - 只扫描边缘行（非全文），降低误判正文重复段落的概率
      - 长度过滤 (2 < len < 80)：排除单字/数字页码与过长的正文行
      - 页数 < 3 时跳过：样本不足，宁可不删也不误删
    """
    if len(raw_pages) < 3:
        return set()

    counts: Counter[str] = Counter()
    for page_text in raw_pages:
        lines = page_text.split("\n")
        edge_lines = lines[:3] + lines[-3:]
        for line in edge_lines:
            stripped = line.strip()
            if 2 < len(stripped) < 80:
                counts[stripped] += 1

    total = len(raw_pages)
    return {line for line, cnt in counts.items() if cnt / total >= threshold}


def _remove_noisy_lines(text: str, noisy: set[str]) -> str:
    """从文本中逐行删除已识别的噪声行。"""
    if not noisy:
        return text
    return "\n".join(ln for ln in text.split("\n") if ln.strip() not in noisy)


# ──────────────────────────────────────────────────────────────
# PDF 解析（含扫描件 OCR 回退）
# ──────────────────────────────────────────────────────────────

# 文本层字符数低于此阈值时触发 OCR（判定为扫描件或图形化页面）
_OCR_TRIGGER_CHARS = 50


def _render_page_to_image(file_bytes: bytes, page_idx: int) -> Image.Image | None:
    """
    用 PyMuPDF 将 PDF 单页渲染为 PIL Image（200 DPI，适合 Tesseract 识别）。
    page_idx: 0-based 页面索引。
    """
    if not _HAS_PYMUPDF:
        return None
    doc = _pymupdf.open(stream=file_bytes, filetype="pdf")
    pix = doc[page_idx].get_pixmap(dpi=200)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)


def _parse_pdf(file_bytes: bytes, source: str) -> list[ParsedPage]:
    """
    PDF 解析主流程：
    1. pypdf 提取文本层（速度快，大多数 PDF 够用）
    2. 页眉/页脚检测并删除
    3. 文本层不足 → PyMuPDF 渲染图像 → Tesseract OCR（chi_sim+eng）
    """
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    total_pages = len(reader.pages)

    # 第一遍：批量提取原始文本（用于噪声行统计）
    raw_texts = [reader.pages[i].extract_text() or "" for i in range(total_pages)]
    noisy = _detect_noisy_lines(raw_texts)
    if noisy:
        logger.debug("检测到页眉/页脚候选行", count=len(noisy), source=source)

    result: list[ParsedPage] = []
    for idx, raw_text in enumerate(raw_texts):
        text = _clean_text(_remove_noisy_lines(raw_text, noisy))

        # 扫描件检测：文本层不足则走 OCR 路径
        if len(text) < _OCR_TRIGGER_CHARS:
            logger.info("触发 OCR 回退", page=idx + 1, source=source)
            img = _render_page_to_image(file_bytes, idx)
            if img is None:
                logger.warning(
                    "PyMuPDF 未安装，OCR 跳过（安装：pip install pymupdf）",
                    page=idx + 1,
                    source=source,
                )
            else:
                ocr_text: str = pytesseract.image_to_string(img, lang="chi_sim+eng")
                text = _clean_text(ocr_text)

        if not text:
            continue

        result.append(ParsedPage(
            content=text,
            page_number=idx + 1,
            # PDF 标题检测依赖字体元数据（pypdf 尚未支持），Phase 1 不实现
            section_title=None,
            source=source,
            metadata={"total_pages": total_pages},
        ))

    return result


# ──────────────────────────────────────────────────────────────
# Word (.docx) 解析
# ──────────────────────────────────────────────────────────────

def _parse_docx(file_bytes: bytes, source: str) -> list[ParsedPage]:
    """
    Word 文档解析：
    - 按 Heading 样式切分章节，每节生成一个 ParsedPage
    - 表格还原为 Markdown 表格格式，保留语义结构
    - 合并单元格去重：相邻列文本相同时只保留一份

    已知限制：
      python-docx 的 doc.paragraphs / doc.tables 分别按顺序遍历，
      无法保证段落与表格的文档内交叉顺序（如"段→表→段"的混排）。
      实际企业文档多为段落集中于节首、表格集中于节末，此限制可接受。
      如需严格保序：改用 doc.element.body.iterchildren() + lxml 解析。
    """
    from docx import Document  # type: ignore[import-untyped]

    doc = Document(io.BytesIO(file_bytes))
    sections: list[ParsedPage] = []
    current_title: str | None = None
    current_lines: list[str] = []
    pseudo_page = 1  # Word 无绝对页码，用章节序号代替

    def _flush() -> None:
        nonlocal pseudo_page
        text = _clean_text("\n".join(current_lines))
        if text:
            sections.append(ParsedPage(
                content=text,
                page_number=pseudo_page,
                section_title=current_title,
                source=source,
                metadata={},
            ))
            pseudo_page += 1

    # 遍历段落：按 Heading 样式识别章节标题
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()
        if not text:
            continue
        if "Heading" in style_name:
            _flush()
            current_lines = []
            current_title = text
        else:
            current_lines.append(text)

    # 遍历表格（追加到当前节末）
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
            # 合并单元格导致相邻列内容重复，逐列去重
            deduped: list[str] = []
            for cell in cells:
                if not deduped or cell != deduped[-1]:
                    deduped.append(cell)
            current_lines.append("| " + " | ".join(deduped) + " |")
            if row_idx == 0:
                current_lines.append("| " + " | ".join(["---"] * len(deduped)) + " |")

    _flush()
    return sections


# ──────────────────────────────────────────────────────────────
# Markdown 解析
# ──────────────────────────────────────────────────────────────

def _parse_markdown(file_bytes: bytes, source: str) -> list[ParsedPage]:
    """
    Markdown 解析：按 H1 / H2 标题切分，每节生成一个 ParsedPage。
    使用 markdown-it-py 解析 token 流，确保标题识别准确（而非简单正则）。
    H3 及以下作为正文处理，不强制断节（避免碎片化）。
    """
    raw_text = file_bytes.decode("utf-8", errors="replace")
    md = MarkdownIt()
    tokens = md.parse(raw_text)

    sections: list[ParsedPage] = []
    current_title: str | None = None
    current_lines: list[str] = []

    def _flush() -> None:
        content = _clean_text("\n".join(current_lines))
        if content:
            sections.append(ParsedPage(
                content=content,
                page_number=None,        # Markdown 无物理页码
                section_title=current_title,
                source=source,
                metadata={},
            ))

    i = 0
    while i < len(tokens):
        tok = tokens[i]
        if tok.type == "heading_open" and tok.tag in ("h1", "h2"):
            heading_text = tokens[i + 1].content.strip() if i + 1 < len(tokens) else ""
            _flush()
            current_lines = []
            current_title = heading_text or None
            i += 3  # heading_open → inline → heading_close
            continue
        if tok.type == "inline" and tok.content:
            current_lines.append(tok.content)
        i += 1

    _flush()
    return sections


# ──────────────────────────────────────────────────────────────
# 纯文本解析
# ──────────────────────────────────────────────────────────────

_TXT_GROUP_SIZE = 10  # 每 N 个段落合并为一个逻辑页（避免碎片化）


def _parse_text(file_bytes: bytes, source: str) -> list[ParsedPage]:
    """纯文本：按双换行切段，每 _TXT_GROUP_SIZE 段合并为一个 ParsedPage。"""
    raw_text = file_bytes.decode("utf-8", errors="replace")
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", raw_text) if p.strip()]

    pages: list[ParsedPage] = []
    for i in range(0, len(paragraphs), _TXT_GROUP_SIZE):
        group = paragraphs[i : i + _TXT_GROUP_SIZE]
        content = _clean_text("\n\n".join(group))
        if content:
            pages.append(ParsedPage(
                content=content,
                page_number=i // _TXT_GROUP_SIZE + 1,
                section_title=None,
                source=source,
                metadata={},
            ))
    return pages


# ──────────────────────────────────────────────────────────────
# 公共接口
# ──────────────────────────────────────────────────────────────

_FORMAT_PARSERS: dict[str, Callable[[bytes, str], list[ParsedPage]]] = {
    ".pdf":      _parse_pdf,
    ".docx":     _parse_docx,
    ".doc":      _parse_docx,        # python-docx 对旧格式有限支持
    ".md":       _parse_markdown,
    ".markdown": _parse_markdown,
    ".txt":      _parse_text,
    ".text":     _parse_text,
}


def parse_document(
    file_bytes: bytes,
    filename: str,
    *,
    format_override: str | None = None,
) -> list[ParsedPage]:
    """
    主入口：按扩展名分发解析器，返回 ParsedPage 列表。

    Args:
        file_bytes:      文件原始字节（调用方负责读取，支持内存和流）
        filename:        原始文件名（含扩展名）；始终作为 ParsedPage.source，保证 citation 溯源正确
        format_override: 强制指定解析格式（扩展名形式，如 ".pdf"）。
                         非 PDF 文档经 LibreOffice 转换后，file_bytes 已是 PDF 字节，
                         但 filename 保留原始名（如 "report.docx"），
                         此时传入 format_override=".pdf" 使分发器走 PDF 解析路径。
                         None（默认）= 按 filename 扩展名自动判断，与原有行为完全一致。

    Returns:
        ParsedPage 列表；空文档返回空列表。
        每条 ParsedPage.source = filename（原始文件名），不受 format_override 影响。

    Raises:
        ValueError: 不支持的文件格式
    """
    # format_override 优先；None 时退回到 filename 扩展名（向后兼容）
    suffix = format_override if format_override is not None else Path(filename).suffix.lower()
    parser_fn = _FORMAT_PARSERS.get(suffix)
    if parser_fn is None:
        supported = ", ".join(_FORMAT_PARSERS)
        raise ValueError(f"不支持的格式 '{suffix}'，当前支持：{supported}")

    log = logger.bind(filename=filename, format=suffix)
    if format_override:
        log = log.bind(format_override=format_override)
    log.info("开始解析文档")
    pages = parser_fn(file_bytes, filename)   # filename 作为 source，不传 format_override
    log.info("文档解析完成", pages=len(pages))
    return pages


def supported_extensions() -> list[str]:
    """返回所有支持的文件扩展名列表。"""
    return list(_FORMAT_PARSERS)
